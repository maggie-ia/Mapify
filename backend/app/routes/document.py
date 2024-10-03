from flask import Blueprint, request, jsonify, send_file, current_app
from werkzeug.utils import secure_filename
from app.models.document import Document
from app.models.user import User
from app import db
from app.utils.file_validators import validate_file_type, validate_file_size
from flask_jwt_extended import jwt_required, get_jwt_identity
from app.services.membership_service import get_page_limit
from app.services.text_processing import (
    summarize_text, paraphrase_text, synthesize_text,
    extract_relevant_phrases, generate_concept_map, translate_text
)
import os
from io import BytesIO
import PyPDF2
import docx

document_bp = Blueprint('document', __name__)

ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

@document_bp.route('/upload', methods=['POST'])
@jwt_required()
def upload_document():
    if 'file' not in request.files:
        return jsonify({"error": "No se ha proporcionado ningún archivo"}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No se ha seleccionado ningún archivo"}), 400
    
    try:
        validate_file_type(file.filename, ALLOWED_EXTENSIONS)
        validate_file_size(file, MAX_FILE_SIZE)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    user_id = get_jwt_identity()
    page_limit = get_page_limit(user_id)
    
    page_count = count_pages(file)
    
    if page_count > page_limit:
        return jsonify({"error": f"El documento excede el límite de {page_limit} páginas para su nivel de membresía"}), 403

    filename = secure_filename(file.filename)
    file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    
    content = extract_text(file_path)
    document = Document(filename=filename, file_path=file_path, user_id=user_id, content=content)
    db.session.add(document)
    db.session.commit()

    return jsonify({"message": "Archivo subido exitosamente", "document_id": document.id}), 201


def count_pages(file):
    file.seek(0)
    if file.filename.endswith('.pdf'):
        reader = PyPDF2.PdfReader(file)
        return len(reader.pages)
    elif file.filename.endswith('.docx'):
        doc = docx.Document(file)
        return len(doc.paragraphs)
    elif file.filename.endswith('.txt'):
        return sum(1 for line in file)
    else:
        return 1  # Default to 1 page for unsupported formats

def extract_text(file_path):
    if file_path.endswith('.pdf'):
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            return " ".join(page.extract_text() for page in reader.pages)
    elif file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        return " ".join(paragraph.text for paragraph in doc.paragraphs)
    elif file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    else:
        return ""  # Return empty string for unsupported formats

@document_bp.route('/<int:document_id>', methods=['GET'])
@jwt_required()
def get_document(document_id):
    document = Document.query.get_or_404(document_id)
    if document.user_id != get_jwt_identity():
        return jsonify({"error": "Acceso no autorizado"}), 403
    return jsonify({"id": document.id, "filename": document.filename, "created_at": document.created_at}), 200

@document_bp.route('/<int:document_id>', methods=['DELETE'])
@jwt_required()
def delete_document(document_id):
    document = Document.query.get_or_404(document_id)
    if document.user_id != get_jwt_identity():
        return jsonify({"error": "Acceso no autorizado"}), 403
    db.session.delete(document)
    db.session.commit()
    return jsonify({"message": "Documento eliminado exitosamente"}), 200

@document_bp.route('/user', methods=['GET'])
@jwt_required()
def get_user_documents():
    user_id = get_jwt_identity()
    documents = Document.query.filter_by(user_id=user_id).all()
    return jsonify([{"id": doc.id, "filename": doc.filename, "created_at": doc.created_at} for doc in documents]), 20

@document_bp.route('/list', methods=['GET'])
@jwt_required()
def list_documents():
    user_id = get_jwt_identity()
    documents = Document.query.filter_by(user_id=user_id).all()
    return jsonify([{
        "id": doc.id,
        "filename": doc.filename,
        "created_at": doc.created_at,
        "file_type": doc.file_type
    } for doc in documents]), 200

@document_bp.route('/<int:doc_id>/summarize', methods=['POST'])
@jwt_required()
def summarize_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    if document.user_id != get_jwt_identity():
        return jsonify({"message": "Unauthorized"}), 403
    
    summary = summarize_text(document.content)
    return jsonify({"summary": summary}), 200

@document_bp.route('/<int:doc_id>/paraphrase', methods=['POST'])
@jwt_required()
def paraphrase_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    if document.user_id != get_jwt_identity():
        return jsonify({"message": "Unauthorized"}), 403
    
    paraphrased = paraphrase_text(document.content)
    return jsonify({"paraphrased": paraphrased}), 200

@document_bp.route('/<int:doc_id>/synthesize', methods=['POST'])
@jwt_required()
def synthesize_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    if document.user_id != get_jwt_identity():
        return jsonify({"message": "Unauthorized"}), 403
    
    synthesis = synthesize_text(document.content)
    return jsonify({"synthesis": synthesis}), 200

@document_bp.route('/<int:doc_id>/relevant_phrases', methods=['POST'])
@jwt_required()
def extract_relevant_phrases_from_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    if document.user_id != get_jwt_identity():
        return jsonify({"message": "Unauthorized"}), 403
    
    relevant_phrases = extract_relevant_phrases(document.content)
    return jsonify({"relevant_phrases": relevant_phrases}), 200

@document_bp.route('/<int:doc_id>/concept_map', methods=['POST'])
@jwt_required()
def create_concept_map(doc_id):
    document = Document.query.get_or_404(doc_id)
    if document.user_id != get_jwt_identity():
        return jsonify({"message": "Unauthorized"}), 403
    
    user = User.query.get(get_jwt_identity())
    max_nodes = 6  # Por defecto para membresía básica
    if user.membership_type == 'premium':
        max_nodes = None  # Sin límite para membresía premium
    
    concept_map_image = generate_concept_map(document.content, max_nodes)
    
    return send_file(
        BytesIO(concept_map_image),
        mimetype='image/png',
        as_attachment=True,
        download_name=f'concept_map_{doc_id}.png'
    )

@document_bp.route('/<int:doc_id>/translate', methods=['POST'])
@jwt_required()
def translate_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    if document.user_id != get_jwt_identity():
        return jsonify({"message": "Unauthorized"}), 403
    
    target_language = request.json.get('target_language')
    if not target_language:
        return jsonify({"message": "Target language is required"}), 400
    
    user = User.query.get(get_jwt_identity())
    allowed_languages = ['en', 'es', 'fr', 'de']  # Idiomas permitidos para membresía básica
    if user.membership_type != 'premium' and target_language not in allowed_languages:
        return jsonify({"message": "Language not available for your membership level"}), 403
    
    translated_text = translate_text(document.content, target_language)
    return jsonify({"translated_text": translated_text}), 200